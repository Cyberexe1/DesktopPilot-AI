"""
AI Meeting Assistant Controller
Captures system audio (WASAPI loopback) → AWS Transcribe → Bedrock summary → .docx
Works with Zoom, Google Meet, Teams, or any meeting tool on Windows.
"""

import asyncio
import json
import logging
import os
import threading
import wave
from datetime import datetime
from pathlib import Path

import boto3

log = logging.getLogger(__name__)

# ── State ──────────────────────────────────────────────────────────────────────
_recording: bool = False
_frames: list    = []
_thread: threading.Thread | None = None
_meeting_title: str = ""
_meeting_start: datetime | None = None
_sample_rate: int = 16000
_channels: int    = 1

REGION     = os.getenv("AWS_DEFAULT_REGION", "us-east-1")
MODEL_ID   = os.getenv("BEDROCK_MODEL_ID", "us.amazon.nova-pro-v1:0")
S3_BUCKET  = os.getenv("S3_BUCKET_NAME", "desktoppilot-audio")
DESKTOP    = Path(os.path.expanduser("~")) / "Desktop"


# ── Audio capture ──────────────────────────────────────────────────────────────

def _capture_audio() -> None:
    """Background thread: captures Windows system audio via WASAPI loopback."""
    global _recording, _frames, _sample_rate, _channels
    try:
        import sounddevice as sd
        import numpy as np

        # Find WASAPI loopback device (captures all system audio output)
        devices = sd.query_devices()
        loopback_idx = None
        for i, dev in enumerate(devices):
            name = dev["name"].lower()
            # Look for a loopback/output device that can be used as input
            if dev["max_input_channels"] > 0 and (
                "loopback" in name or
                "stereo mix" in name or
                "what u hear" in name or
                "wave out mix" in name
            ):
                loopback_idx = i
                log.info(f"Found loopback device: {dev['name']}")
                break

        if loopback_idx is None:
            # Fallback: use default microphone
            log.warning("No loopback device found — falling back to microphone input")
            loopback_idx = None  # sd.rec uses default

        _sample_rate = 16000
        _channels = 1

        with sd.InputStream(
            samplerate=_sample_rate,
            channels=_channels,
            dtype="int16",
            device=loopback_idx,
            blocksize=1024,
        ) as stream:
            log.info("Audio capture started")
            while _recording:
                data, _ = stream.read(1024)
                _frames.append(data.tobytes())

    except Exception as e:
        log.error(f"Audio capture error: {e}")


# ── Public API ─────────────────────────────────────────────────────────────────

def start_meeting(title: str = "") -> str:
    """Start recording system audio for a meeting."""
    global _recording, _frames, _thread, _meeting_title, _meeting_start

    if _recording:
        return "Meeting recording is already in progress"

    _meeting_title = title or f"Meeting_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
    _meeting_start = datetime.now()
    _frames = []
    _recording = True

    _thread = threading.Thread(target=_capture_audio, daemon=True)
    _thread.start()

    from controllers.voice_output_controller import speak
    speak(f"Meeting recording started. Say stop meeting when you're done.")
    log.info(f"Meeting started: {_meeting_title}")
    return f"Recording started — '{_meeting_title}'. Say 'stop meeting' when done."


def stop_meeting() -> dict:
    """Stop recording and save audio to a WAV file."""
    global _recording, _frames, _thread

    if not _recording:
        return {"error": "No meeting recording in progress"}

    _recording = False
    if _thread:
        _thread.join(timeout=3)

    if not _frames:
        return {"error": "No audio captured. Check your microphone or loopback device."}

    # Save WAV file to Desktop
    ts       = datetime.now().strftime("%Y%m%d_%H%M%S")
    wav_name = f"meeting_{ts}.wav"
    wav_path = str(DESKTOP / wav_name)

    try:
        with wave.open(wav_path, "wb") as wf:
            wf.setnchannels(_channels)
            wf.setsampwidth(2)  # 16-bit = 2 bytes
            wf.setframerate(_sample_rate)
            for chunk in _frames:
                wf.writeframes(chunk)
        log.info(f"Meeting audio saved: {wav_path} ({len(_frames)} chunks)")
    except Exception as e:
        return {"error": f"Failed to save audio: {e}"}

    return {
        "wav_path": wav_path,
        "title":    _meeting_title,
        "duration": str(datetime.now() - _meeting_start).split(".")[0],
    }


def get_meeting_status() -> dict:
    return {
        "recording": _recording,
        "title":     _meeting_title,
        "started_at": _meeting_start.isoformat() if _meeting_start else None,
        "duration":   str(datetime.now() - _meeting_start).split(".")[0] if _meeting_start else None,
    }


# ── Transcription ──────────────────────────────────────────────────────────────

async def _transcribe_wav(wav_path: str, title: str) -> str:
    """Upload WAV to S3 and transcribe with AWS Transcribe (speaker diarization)."""
    loop = asyncio.get_event_loop()
    return await loop.run_in_executor(None, _transcribe_sync, wav_path, title)


def _transcribe_sync(wav_path: str, title: str) -> str:
    s3  = boto3.client("s3",         region_name=REGION)
    txc = boto3.client("transcribe", region_name=REGION)

    # Upload audio to S3
    s3_key = f"meetings/{Path(wav_path).name}"
    log.info(f"Uploading {wav_path} → s3://{S3_BUCKET}/{s3_key}")
    s3.upload_file(wav_path, S3_BUCKET, s3_key)

    s3_uri = f"s3://{S3_BUCKET}/{s3_key}"
    job_name = f"meeting_{datetime.now().strftime('%Y%m%d%H%M%S')}"

    log.info(f"Starting Transcribe job: {job_name}")
    txc.start_transcription_job(
        TranscriptionJobName=job_name,
        Media={"MediaFileUri": s3_uri},
        MediaFormat="wav",
        LanguageCode="en-US",
        Settings={
            "ShowSpeakerLabels": True,
            "MaxSpeakerLabels": 10,   # support up to 10 meeting participants
        },
    )

    # Poll until complete
    import time
    for _ in range(120):  # max 4 minutes
        time.sleep(2)
        resp   = txc.get_transcription_job(TranscriptionJobName=job_name)
        status = resp["TranscriptionJob"]["TranscriptionJobStatus"]
        if status == "COMPLETED":
            break
        if status == "FAILED":
            reason = resp["TranscriptionJob"].get("FailureReason", "Unknown")
            raise RuntimeError(f"Transcribe job failed: {reason}")

    # Download transcript JSON
    import urllib.request
    transcript_uri = resp["TranscriptionJob"]["Transcript"]["TranscriptFileUri"]
    with urllib.request.urlopen(transcript_uri) as f:
        data = json.load(f)

    return _format_transcript(data)


def _format_transcript(data: dict) -> str:
    """Format diarized transcript as 'Speaker X [HH:MM:SS]: text'."""
    items    = data.get("results", {}).get("items", [])
    segments = data.get("results", {}).get("speaker_labels", {}).get("segments", [])

    # Build time → speaker map
    speaker_map: dict[str, str] = {}
    for seg in segments:
        speaker = seg.get("speaker_label", "Speaker")
        for item in seg.get("items", []):
            start = item.get("start_time", "")
            if start:
                speaker_map[start] = speaker

    lines: list[str] = []
    current_speaker   = None
    current_words: list[str] = []
    current_start    = None

    for item in items:
        if item["type"] == "punctuation":
            if current_words:
                current_words[-1] += item["alternatives"][0]["content"]
            continue

        start   = item.get("start_time", "")
        speaker = speaker_map.get(start, current_speaker or "Speaker 1")
        word    = item["alternatives"][0]["content"]

        if speaker != current_speaker:
            if current_words and current_speaker:
                ts = _fmt_time(current_start)
                lines.append(f"[{current_speaker} — {ts}]: {' '.join(current_words)}")
            current_speaker = speaker
            current_words   = [word]
            current_start   = start
        else:
            current_words.append(word)

    if current_words and current_speaker:
        ts = _fmt_time(current_start)
        lines.append(f"[{current_speaker} — {ts}]: {' '.join(current_words)}")

    return "\n".join(lines) if lines else data.get("results", {}).get("transcripts", [{}])[0].get("transcript", "")


def _fmt_time(seconds_str: str) -> str:
    try:
        secs = float(seconds_str)
        m, s = divmod(int(secs), 60)
        h, m = divmod(m, 60)
        return f"{h:02d}:{m:02d}:{s:02d}"
    except Exception:
        return "00:00:00"


# ── Bedrock Summary ────────────────────────────────────────────────────────────

def _summarize_sync(transcript: str, title: str) -> dict:
    """Send transcript to Bedrock Nova Pro → structured summary JSON."""
    bedrock = boto3.client("bedrock-runtime", region_name=REGION)

    prompt = f"""You are an expert meeting analyst. Analyze the following meeting transcript and produce a structured JSON summary.

Meeting title: {title}

Transcript:
{transcript[:12000]}

Return ONLY valid JSON in this exact format:
{{
  "summary": "2-3 sentence overview of the entire meeting",
  "key_topics": ["topic 1", "topic 2", "topic 3"],
  "action_items": [
    {{"speaker": "Speaker 1", "action": "description of what they need to do", "deadline": "by Friday / ASAP / no deadline"}}
  ],
  "decisions": ["decision 1", "decision 2"],
  "participants": 2,
  "sentiment": "positive / neutral / tense",
  "follow_up_needed": true
}}"""

    is_nova = "nova" in MODEL_ID.lower() or "amazon" in MODEL_ID.lower()
    if is_nova:
        body = {
            "schemaVersion": "messages-v1",
            "messages": [{"role": "user", "content": [{"text": prompt}]}],
            "inferenceConfig": {"maxTokens": 1500, "temperature": 0.1},
        }
    else:
        body = {
            "prompt": prompt,
            "max_gen_len": 1500,
            "temperature": 0.1,
        }

    resp = bedrock.invoke_model(
        modelId=MODEL_ID,
        body=json.dumps(body),
        contentType="application/json",
        accept="application/json",
    )
    raw = json.loads(resp["body"].read())

    # Extract text from response
    if "output" in raw and "message" in raw["output"]:
        text = raw["output"]["message"]["content"][0]["text"]
    elif "generation" in raw:
        text = raw["generation"]
    else:
        text = str(raw)

    # Parse JSON from the response
    import re
    match = re.search(r'\{.*\}', text, re.DOTALL)
    if match:
        try:
            return json.loads(match.group())
        except Exception:
            pass

    # Fallback minimal summary
    return {
        "summary": text[:500],
        "key_topics": [],
        "action_items": [],
        "decisions": [],
        "participants": 1,
        "sentiment": "neutral",
        "follow_up_needed": False,
    }


# ── Word Doc generation ────────────────────────────────────────────────────────

def _build_docx(title: str, date_str: str, duration: str,
                transcript: str, summary: dict) -> str:
    """Build a formatted .docx meeting notes file on the Desktop."""
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # ── Title ──
    heading = doc.add_heading(f"Meeting Notes: {title}", level=0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = heading.runs[0]
    run.font.color.rgb = RGBColor(0xCC, 0x22, 0x00)  # brand red

    # ── Meta ──
    doc.add_paragraph(f"Date: {date_str}   |   Duration: {duration}")
    doc.add_paragraph(f"Participants detected: {summary.get('participants', '?')}   |   Sentiment: {summary.get('sentiment', 'neutral').title()}")
    doc.add_paragraph("")

    # ── Summary ──
    doc.add_heading("Executive Summary", level=1)
    doc.add_paragraph(summary.get("summary", ""))

    # ── Key topics ──
    topics = summary.get("key_topics", [])
    if topics:
        doc.add_heading("Key Topics Discussed", level=1)
        for t in topics:
            doc.add_paragraph(f"• {t}")

    # ── Action items ──
    actions = summary.get("action_items", [])
    if actions:
        doc.add_heading("Action Items", level=1)
        for item in actions:
            speaker  = item.get("speaker", "Unassigned")
            action   = item.get("action", "")
            deadline = item.get("deadline", "")
            line = f"[ ] {speaker}: {action}"
            if deadline and deadline.lower() not in ("no deadline", "none", ""):
                line += f"  — Due: {deadline}"
            p = doc.add_paragraph(line)
            p.runs[0].font.bold = True

    # ── Decisions ──
    decisions = summary.get("decisions", [])
    if decisions:
        doc.add_heading("Decisions Made", level=1)
        for d in decisions:
            doc.add_paragraph(f"✓ {d}")

    # ── Full transcript ──
    doc.add_heading("Full Transcript", level=1)
    doc.add_paragraph("(Speaker labels: Speaker 1, Speaker 2, ... correspond to individual meeting participants)")
    doc.add_paragraph("")

    for line in transcript.split("\n"):
        if line.strip():
            p = doc.add_paragraph()
            if line.startswith("[Speaker"):
                # Speaker label — bold
                colon_idx = line.find("]:")
                if colon_idx != -1:
                    label = line[:colon_idx + 2]
                    text  = line[colon_idx + 2:]
                    run1  = p.add_run(label + " ")
                    run1.bold = True
                    run1.font.color.rgb = RGBColor(0xCC, 0x22, 0x00)
                    p.add_run(text)
                else:
                    p.add_run(line)
            else:
                p.add_run(line)

    # Save
    ts       = datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = f"MeetingNotes_{title.replace(' ', '_')}_{ts}.docx"
    doc_path = str(DESKTOP / filename)
    doc.save(doc_path)
    log.info(f"Meeting notes saved: {doc_path}")
    return doc_path


# ── Main orchestrator ──────────────────────────────────────────────────────────

async def process_meeting(wav_path: str, title: str, send_email_to: str = "") -> dict:
    """
    Full pipeline:
      WAV → S3 → AWS Transcribe → Bedrock summary → .docx → (optional email)
    """
    from controllers.voice_output_controller import speak

    speak("Processing your meeting. Transcribing audio with AWS...")
    log.info(f"Processing meeting: {title} | wav: {wav_path}")

    # 1. Transcribe
    try:
        transcript = await _transcribe_wav(wav_path, title)
        log.info(f"Transcript length: {len(transcript)} chars")
    except Exception as e:
        log.error(f"Transcription failed: {e}")
        return {"error": f"Transcription failed: {e}"}

    # 2. Summarize with Bedrock
    speak("Generating meeting summary with Amazon Bedrock...")
    loop = asyncio.get_event_loop()
    try:
        summary = await loop.run_in_executor(None, _summarize_sync, transcript, title)
    except Exception as e:
        log.error(f"Summary failed: {e}")
        summary = {
            "summary": "Summary generation failed.",
            "key_topics": [], "action_items": [],
            "decisions": [], "participants": 1,
            "sentiment": "neutral", "follow_up_needed": False,
        }

    # 3. Build Word document
    date_str = _meeting_start.strftime("%B %d, %Y %I:%M %p") if _meeting_start else datetime.now().strftime("%B %d, %Y")
    duration = str(datetime.now() - _meeting_start).split(".")[0] if _meeting_start else "Unknown"

    try:
        doc_path = await loop.run_in_executor(
            None, _build_docx, title, date_str, duration, transcript, summary
        )
    except Exception as e:
        log.error(f"Doc build failed: {e}")
        doc_path = ""

    # 4. Open the document
    if doc_path and os.path.exists(doc_path):
        try:
            os.startfile(doc_path)
        except Exception:
            pass

    # 5. Optional email
    if send_email_to:
        try:
            subject = f"Meeting Notes: {title} — {date_str}"
            body    = (
                f"Hi,\n\nPlease find the meeting notes for '{title}' held on {date_str}.\n\n"
                f"SUMMARY:\n{summary.get('summary', '')}\n\n"
                f"ACTION ITEMS:\n" +
                "\n".join(
                    f"• {a['speaker']}: {a['action']}" + (f" (Due: {a['deadline']})" if a.get('deadline') else "")
                    for a in summary.get("action_items", [])
                ) +
                f"\n\nFull notes document is attached.\n\nRegards,\nDesktopPilot AI"
            )
            from controllers.browser_controller import open_gmail_compose
            await open_gmail_compose(send_email_to, subject, body)
            log.info(f"Meeting summary email opened for: {send_email_to}")
        except Exception as e:
            log.warning(f"Email open failed: {e}")

    n_actions = len(summary.get("action_items", []))
    n_topics  = len(summary.get("key_topics", []))
    speak(
        f"Meeting notes ready. {n_actions} action items and {n_topics} key topics extracted. "
        f"Document saved to your Desktop."
    )

    return {
        "title":        title,
        "duration":     duration,
        "doc_path":     doc_path,
        "transcript":   transcript,
        "summary":      summary.get("summary", ""),
        "key_topics":   summary.get("key_topics", []),
        "action_items": summary.get("action_items", []),
        "decisions":    summary.get("decisions", []),
        "participants": summary.get("participants", 1),
        "sentiment":    summary.get("sentiment", "neutral"),
    }
